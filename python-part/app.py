from flask import Flask, request, jsonify, render_template
from flask_cors import CORS
import google.generativeai as genai
import os
from dotenv import load_dotenv
import pdfplumber
import docx
from io import BytesIO
import base64

# Load environment variables
load_dotenv()

# Configure Gemini model
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))
model = genai.GenerativeModel("gemini-2.0-flash")

# Initialize Flask app
app = Flask(__name__)
CORS(app)

@app.route('/')
def home():
    return render_template('landing.html')

@app.route('/index')
def index():
    return render_template('index.html')  # This will load index.html


@app.route('/analyzecode')
def codelyzer():
    return render_template('code_analyzer.html')

@app.route('/codegen')
def codegena():
    return render_template('code_builder.html')

@app.route('/debug')
def debugger():
    return render_template('code_debugger.html')


@app.route('/mmap')
def mmap():
    return render_template('mmap.html')
@app.route('/edit')
def edit():
    
    return render_template('edit.html')
@app.route('/mme')
def mme():
    return render_template('mind-map-editing.html')

@app.route('/new')
def new():
    query1 = request.args.get('q', '')  # Get the `q` query param
    print(query1)
    return render_template('new.html', query=query1)


@app.route('/summarize', methods=['POST'])
def summarize():
    data = request.json
    input_type = data.get('input_type')
    content = data.get('content')

    try:
        if input_type == 'text':
            text = content
        elif input_type == 'pdf':
            file_bytes = BytesIO(base64.b64decode(content))
            with pdfplumber.open(file_bytes) as pdf:
                text = "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
        elif input_type == 'docx':
            file_bytes = BytesIO(base64.b64decode(content))
            doc = docx.Document(file_bytes)
            text = "\n".join([para.text for para in doc.paragraphs])
        else:
            return jsonify({'error': 'Unsupported input type'}), 400

        prompt = f"Summarize the following notes:\n\n{text}"
        response = model.generate_content(prompt)
        return jsonify({'summary': response.text})

    except Exception as e:
        return jsonify({'error': str(e)}), 500
    
@app.route('/codelyzer', methods=['POST'])
def codelyze():
    data = request.json
    input_type = data.get('input_type')
    content = data.get('content')

    try:
        if input_type == 'text':
            text = content
        elif input_type == 'pdf':
            file_bytes = BytesIO(base64.b64decode(content))
            with pdfplumber.open(file_bytes) as pdf:
                text = "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
        elif input_type == 'docx':
            file_bytes = BytesIO(base64.b64decode(content))
            doc = docx.Document(file_bytes)
            text = "\n".join([para.text for para in doc.paragraphs])
        else:
            return jsonify({'error': 'Unsupported input type'}), 400

        prompt = f"Analyze the following code for clarity, performance, security, and best practices. Identify any bugs, anti-patterns, or areas for improvement. Summarize what the code does and suggest optimizations or refactoring opportunities. Include recommendations in bullet points.\n\n{text}"
        response = model.generate_content(prompt)
        print(response.text)
        return jsonify({'summary': response.text})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/codegena', methods=['POST'])
def codegen():
    data = request.json
    input_type = data.get('input_type')
    content = data.get('content')

    try:
        if input_type == 'text':
            text = content
        elif input_type == 'pdf':
            file_bytes = BytesIO(base64.b64decode(content))
            with pdfplumber.open(file_bytes) as pdf:
                text = "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
        elif input_type == 'docx':
            file_bytes = BytesIO(base64.b64decode(content))
            doc = docx.Document(file_bytes)
            text = "\n".join([para.text for para in doc.paragraphs])
        else:
            return jsonify({'error': 'Unsupported input type'}), 400

        prompt = f"given the following conditions provide a neat concise and meaningful code covering all the corner cases.\n\n{text}"
        response = model.generate_content(prompt)
        print(response.text)
        return jsonify({'summary': response.text})

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/debugger', methods=['POST'])
def debugge():
    data = request.json
    input_type = data.get('input_type')
    content = data.get('content')
    try:
        if input_type == 'text':
            text = content
        elif input_type == 'pdf':
            file_bytes = BytesIO(base64.b64decode(content))
            with pdfplumber.open(file_bytes) as pdf:
                text = "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
        elif input_type == 'docx':
            file_bytes = BytesIO(base64.b64decode(content))
            doc = docx.Document(file_bytes)
            text = "\n".join([para.text for para in doc.paragraphs])
        else:
            return jsonify({'error': 'Unsupported input type'}), 400

        prompt = f"given the following code check for all the corner cases and wherever it might break down.\n\n{text}"
        response = model.generate_content(prompt)
        print(response.text)
        return jsonify({'summary': response.text})

    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True, port=3002)

